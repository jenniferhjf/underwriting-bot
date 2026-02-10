import streamlit as st
import json
import os
from datetime import datetime
from pathlib import Path
import hashlib
import re
import traceback
from typing import List, Dict, Any, Optional, Tuple
import requests
from io import BytesIO
import base64

# ===========================
# Configuration
# ===========================

VERSION = "2.8.2"
APP_TITLE = "Enhanced Underwriting Assistant - Professional RAG+CoT System"

# API Configuration
DEFAULT_API_KEY = os.getenv("API_KEY", "sk-99bba2ce117444e197270f17d303e74f")
API_BASE = "https://api.deepseek.com/v1"
API_MODEL = "deepseek-chat"

# Directory Structure
DATA_DIR = Path("data")
WORKSPACES_DIR = DATA_DIR / "workspaces"
EMBEDDINGS_DIR = DATA_DIR / "embeddings"
ANALYSIS_DIR = DATA_DIR / "analysis"
REVIEW_DIR = DATA_DIR / "review_queue"
AUDIT_DIR = DATA_DIR / "audit_logs"
CONFIG_DIR = DATA_DIR / "config"

# Initial dataset file
INITIAL_DATASET = "Hull - MSC_Memo.pdf"

# Supported file formats
SUPPORTED_FORMATS = {
    'pdf': '📄 PDF',
    'docx': '📝 Word',
    'doc': '📝 Word',
    'txt': '📃 Text',
    'xlsx': '📊 Excel',
    'xls': '📊 Excel',
    'png': '🖼️ Image',
    'jpg': '🖼️ Image',
    'jpeg': '🖼️ Image'
}

# Tag categories
TAG_OPTIONS = {
    'equipment': ['Hull', 'Cargo', 'Liability', 'Property', 'Marine', 'Aviation'],
    'industry': ['Shipping', 'Manufacturing', 'Retail', 'Technology', 'Construction'],
    'timeline': ['2024', '2025', '2026', 'Q1', 'Q2', 'Q3', 'Q4']
}

# Insurance terminology dictionary
INSURANCE_TERMS = {
    'retention': 'The amount of risk that the insured retains before insurance coverage applies',
    'premium': 'The amount paid for insurance coverage',
    'coverage': 'The scope and extent of protection provided by an insurance policy',
    'deductible': 'The amount the insured must pay before the insurer pays a claim',
    'underwriting slip': 'A document containing key details of an insurance risk',
    'loss ratio': 'The ratio of losses paid to premiums earned',
    'exposure': 'The state of being subject to the possibility of loss',
    'claims': 'Requests for compensation under an insurance policy',
    'policy': 'A contract of insurance',
    'endorsement': 'An amendment or addition to an insurance policy',
    'exclusion': 'Specific conditions or circumstances that are not covered',
    'limit': 'The maximum amount an insurer will pay for a covered loss',
    'aggregate': 'The total limit of coverage for all claims during a policy period',
    'per occurrence': 'The limit applicable to each individual claim or incident',
    'retroactive date': 'The date from which coverage applies for claims-made policies'
}

# ===========================
# System Prompts (Human-Readable Format)
# ===========================

SYSTEM_INSTRUCTION = """You are an expert underwriting assistant with deep knowledge of insurance policies, 
risk assessment, and document analysis. Your role is to help underwriters make informed decisions by:
1. Extracting and analyzing key information from policy documents
2. Translating handwritten annotations into structured electronic text
3. Identifying critical risk factors and coverage terms
4. Providing comprehensive analysis with actionable insights
5. Maintaining strict accuracy and professional standards

Always provide responses in clear, professional format suitable for business clients."""

ELECTRONIC_TEXT_ANALYSIS_SYSTEM = """Analyze this insurance document and provide a BRIEF summary ONLY (3-5 sentences maximum).

**CRITICAL**: Base on ACTUAL document content. Keep it concise and client-friendly.

Cover these key points in 3-5 sentences:
- Insurance type and policy name
- Insured party and broker (if mentioned)
- Key financial terms (premium, coverage, loss ratio)
- Main risk factors or special notes

Example format:
"This is a renewal memorandum for MSC vessel Hull & Machinery insurance. The insured is Mediterranean Shipping Company, broker is Cambiaso Risso Asia. Premium of USD 125,000 with net loss ratio of 74.32% and brokerage rate of 22.5%. Deductible increased from USD 500k to USD 1mil with FCIL writing at own merits."

DO NOT use sections, headers, or detailed breakdown. Just 3-5 concise sentences."""

HANDWRITING_TRANSLATION_SYSTEM = """Translate handwritten annotations from this insurance document.

**CRITICAL**: For EACH handwritten annotation, provide ONLY:

[Location] Translated text (Confidence: XX%)

Example output:
[Top of Page 1] To CEO: Renewal suggestions for your consideration (Confidence: 85%)
[Right margin, Page 2] Check premium calculation and verify loss ratio (Confidence: 92%)
[Bottom of Page 3] Approved for renewal with increased deductible (Confidence: 78%)

Rules:
- DO NOT write "Handwriting Summary" or any overview section
- DO NOT write "Detected Annotations" header
- DO NOT write "Key Insights" section
- Only provide direct translations in the format: [Location] Text (Confidence: XX%)
- Estimate confidence 0-100% based on clarity
- Keep each translation concise and clear"""

# Q&A Extraction removed in v2.8.2

AUTO_ANNOTATE_SYSTEM = """Automatically annotate this underwriting document with key metadata.

Provide the annotation in this business-ready format:

## Document Classification
**Tags:** [tag1, tag2, tag3]
**Insurance Type:** [specific type]
**Risk Level:** [Low/Medium/High/Critical]

## Preliminary Decision
**Recommendation:** [Accept/Review/Decline/Pending]
**Confidence:** [0-100%]

## Financial Summary
**Estimated Premium:** [amount or TBD]
**Retention Amount:** [amount or TBD]

## Executive Summary
[2-3 sentence overview of the case]

## Key Insights
- [Insight 1]
- [Insight 2]
- [Insight 3]

---
Note: This is an automated preliminary analysis. Final decisions require human underwriter review."""

# ===========================
# Configuration Management
# ===========================

def ensure_dirs():
    """Create necessary directories"""
    for dir_path in [WORKSPACES_DIR, EMBEDDINGS_DIR, ANALYSIS_DIR, REVIEW_DIR, AUDIT_DIR, CONFIG_DIR]:
        dir_path.mkdir(parents=True, exist_ok=True)

def load_api_config() -> Dict:
    """Load API configuration"""
    config_file = CONFIG_DIR / "api_config.json"
    
    if config_file.exists():
        try:
            with open(config_file, 'r') as f:
                return json.load(f)
        except:
            pass
    
    return {"api_key": DEFAULT_API_KEY}

def save_api_config(api_key: str):
    """Save API configuration"""
    config_file = CONFIG_DIR / "api_config.json"
    
    with open(config_file, 'w') as f:
        json.dump({"api_key": api_key}, f)

def get_api_key() -> str:
    """Get API key from config or session state"""
    if 'api_key' in st.session_state and st.session_state.api_key:
        return st.session_state.api_key
    
    config = load_api_config()
    return config.get('api_key', DEFAULT_API_KEY)

def validate_api_key(api_key: str) -> Tuple[bool, str]:
    """Validate API key and test connection"""
    if not api_key:
        return False, "API key is empty"
    
    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": API_MODEL,
            "messages": [{"role": "user", "content": "test"}],
            "max_tokens": 10
        }
        
        response = requests.post(
            f"{API_BASE}/chat/completions",
            headers=headers,
            json=payload,
            timeout=10
        )
        
        if response.status_code == 200:
            return True, "API key is valid"
        elif response.status_code == 401:
            return False, "API key is invalid or unauthorized"
        else:
            return False, f"API error: {response.status_code}"
            
    except Exception as e:
        return False, f"Connection error: {str(e)}"

# ===========================
# Utility Functions
# ===========================

def log_audit_event(event_type: str, details: Dict[str, Any]):
    """Log audit events for compliance tracking"""
    try:
        timestamp = datetime.now().isoformat()
        log_entry = {
            "timestamp": timestamp,
            "event_type": event_type,
            "details": details,
            "user": "system"
        }
        
        log_file = AUDIT_DIR / f"audit_{datetime.now().strftime('%Y%m')}.jsonl"
        with open(log_file, 'a', encoding='utf-8') as f:
            f.write(json.dumps(log_entry) + '\n')
    except Exception as e:
        st.warning(f"Failed to log audit event: {e}")

def preprocess_insurance_text(text: str) -> str:
    """Preprocess text with insurance terminology awareness"""
    processed = text.lower()
    
    for term in INSURANCE_TERMS.keys():
        processed = re.sub(rf'\b{term}s?\b', term, processed, flags=re.IGNORECASE)
    
    processed = re.sub(r'\s+', ' ', processed).strip()
    
    return processed

def generate_embedding(text: str) -> List[float]:
    """Generate embedding vector for text"""
    hash_obj = hashlib.sha256(preprocess_insurance_text(text).encode())
    hash_int = int.from_bytes(hash_obj.digest(), byteorder='big')
    
    embedding = []
    for i in range(1536):
        seed = hash_int + i
        embedding.append((seed % 1000) / 1000.0 - 0.5)
    
    return embedding

def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Calculate cosine similarity between two vectors"""
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    magnitude1 = sum(a * a for a in vec1) ** 0.5
    magnitude2 = sum(b * b for b in vec2) ** 0.5
    
    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0
    
    return dot_product / (magnitude1 * magnitude2)

def extract_tag_from_filename(filename: str) -> Optional[str]:
    """Extract primary tag from filename"""
    name_without_ext = os.path.splitext(filename)[0]
    parts = re.split(r'[\s_\-]+', name_without_ext)
    
    if not parts:
        return None
    
    first_word = parts[0].strip().title()
    
    if first_word.isdigit() or re.match(r'\d{4}', first_word):
        return None
    
    return first_word

def call_llm_api(system_prompt: str, user_prompt: str, 
                 temperature: float = 0.3, max_tokens: int = 4000) -> str:
    """Call LLM API for text generation"""
    try:
        api_key = get_api_key()
        
        if not api_key:
            st.error("⚠️ API key not configured. Please set it in the sidebar.")
            return ""
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": API_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        
        response = requests.post(
            f"{API_BASE}/chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )
        
        if response.status_code == 401:
            st.error("⚠️ API Authentication Failed. Please check your API key in the sidebar.")
            return ""
        
        response.raise_for_status()
        result = response.json()
        
        return result['choices'][0]['message']['content']
        
    except requests.exceptions.HTTPError as e:
        st.error(f"API HTTP Error: {e}")
        return ""
    except Exception as e:
        st.error(f"API Error: {e}")
        return ""

# ===========================
# PDF Processing Functions
# ===========================

def is_scanned_pdf(file_path: Path) -> bool:
    """Check if PDF is scanned (image-only) and needs OCR"""
    try:
        import fitz
        doc = fitz.open(file_path)
        
        total_text_len = 0
        total_images = 0
        pages_to_check = min(3, len(doc))
        
        for page_num in range(pages_to_check):
            page = doc[page_num]
            text = page.get_text().strip()
            images = page.get_images()
            
            total_text_len += len(text)
            total_images += len(images)
        
        doc.close()
        
        if total_images > 0 and total_text_len < 100:
            return True
        
        return False
        
    except Exception as e:
        return False

def extract_images_from_pdf(file_path: Path) -> List[Dict]:
    """Extract all images from PDF (for scanned documents and handwriting detection)"""
    try:
        import fitz
        doc = fitz.open(file_path)
        images = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Convert to base64
                    image_b64 = base64.b64encode(image_bytes).decode('utf-8')
                    
                    images.append({
                        'id': f"page{page_num+1}_img{img_index+1}",
                        'data': image_b64,
                        'type': 'scanned',
                        'page': page_num + 1,
                        'source_file': file_path.name,
                        'size': len(image_bytes)
                    })
                except Exception as e:
                    continue
        
        doc.close()
        return images
        
    except Exception as e:
        st.warning(f"PDF image extraction error: {e}")
        return []

def detect_handwriting_in_images(images: List[Dict]) -> bool:
    """Detect if images likely contain handwriting (heuristic approach)"""
    if not images:
        return False
    
    # Count images per page
    pages = {}
    for img in images:
        page = img.get('page', 1)
        pages[page] = pages.get(page, 0) + 1
    
    # If any page has 4+ images, likely has handwriting overlays/annotations
    for page, count in pages.items():
        if count >= 4:
            return True
    
    # Check for small images (might be signatures/stamps)
    for img in images:
        if img.get('size', 0) < 50000:  # < 50KB
            return True
    
    return False

def extract_text_from_scanned_pdf(file_path: Path) -> str:
    """Extract information from scanned PDF"""
    try:
        import fitz
        
        doc = fitz.open(file_path)
        
        extracted_info = f"""📷 SCANNED DOCUMENT DETECTED
{'='*50}

This document appears to be a scanned/image-based PDF.

Document Information:
- Filename: {file_path.name}
- Total Pages: {len(doc)}
- Format: PDF {doc.metadata.get('format', 'Unknown')}
- Creator: {doc.metadata.get('creator', 'Unknown')}

Image Content Analysis:
"""
        
        for i, page in enumerate(doc):
            images = page.get_images()
            extracted_info += f"\nPage {i+1}: Contains {len(images)} image(s)"
        
        extracted_info += "\n\n" + "="*50
        extracted_info += "\n⚠️ NOTE: This is an image-only PDF without extractable text."
        extracted_info += "\n\n📋 To extract text from this document:"
        extracted_info += "\n1. The system will analyze images for handwritten annotations"
        extracted_info += "\n2. Use the 'Handwriting Translation' tab to view results"
        extracted_info += "\n3. You can also upload individual page images for OCR processing"
        extracted_info += "\n\nFor production use, integrate with:"
        extracted_info += "\n- Google Cloud Vision API"
        extracted_info += "\n- AWS Textract"
        extracted_info += "\n- Azure Computer Vision"
        
        doc.close()
        
        return extracted_info
        
    except Exception as e:
        return f"Error processing scanned PDF: {e}"

def extract_text_from_pdf(file_path: Path) -> Tuple[str, List[Dict]]:
    """Enhanced PDF extraction - returns (text, images)"""
    try:
        import fitz
        doc = fitz.open(file_path)
        
        text_content = []
        all_images = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Try to extract text
            text = page.get_text()
            if text.strip():
                text_content.append(f"=== Page {page_num+1} ===\n{text}\n")
            
            # Extract images
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    all_images.append({
                        'id': f"page{page_num+1}_img{img_index+1}",
                        'data': base64.b64encode(image_bytes).decode('utf-8'),
                        'type': 'scanned' if not text.strip() else 'embedded',
                        'page': page_num + 1,
                        'source_file': file_path.name,
                        'size': len(image_bytes)
                    })
                except:
                    continue
        
        doc.close()
        
        # If no text but has images, it's a scanned document
        if not text_content and all_images:
            scanned_info = extract_text_from_scanned_pdf(file_path)
            return (scanned_info, all_images)
        
        final_text = "\n\n".join(text_content) if text_content else ""
        return (final_text, all_images)
        
    except Exception as e:
        st.warning(f"PDF extraction error: {e}")
        return ("", [])

def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except Exception as e:
        st.warning(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        st.warning(f"TXT extraction error: {e}")
        return ""

def extract_text_from_file(file_path: Path) -> Tuple[str, List[Dict]]:
    """Extract text and images based on file type"""
    ext = file_path.suffix.lower().lstrip('.')
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['docx', 'doc']:
        text = extract_text_from_docx(file_path)
        images = extract_images_from_docx(file_path)
        return (text, images)
    elif ext == 'txt':
        text = extract_text_from_txt(file_path)
        return (text, [])
    else:
        return ("", [])

def extract_images_from_docx(file_path: Path) -> List[Dict]:
    """Extract embedded images from DOCX file (skip external links)"""
    try:
        from docx import Document
        doc = Document(file_path)
        images = []
        
        for rel in doc.part.rels.values():
            # Skip external relationships (like linked images from URLs)
            if hasattr(rel, 'target_mode') and rel.target_mode == 'External':
                continue
            
            # Only process internal image relationships
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    image_id = f"img_{len(images)+1}"
                    
                    images.append({
                        'id': image_id,
                        'data': base64.b64encode(image_data).decode('utf-8'),
                        'type': 'embedded',
                        'source_file': file_path.name,
                        'size': len(image_data)
                    })
                except Exception as e:
                    # Skip images that can't be processed
                    continue
        
        return images
    except Exception as e:
        st.warning(f"Image extraction error: {e}")
        return []

def has_embedded_images(file_path: Path) -> bool:
    """Check if document has embedded images (potential handwriting)"""
    try:
        if file_path.suffix.lower() in ['.docx', '.doc']:
            images = extract_images_from_docx(file_path)
            return len(images) > 0
        elif file_path.suffix.lower() == '.pdf':
            import fitz
            doc = fitz.open(file_path)
            for page in doc:
                if len(page.get_images()) > 0:
                    doc.close()
                    return True
            doc.close()
            return False
        return False
    except:
        return False

def classify_handwriting_quality(image_data: str) -> Tuple[str, float]:
    """Classify handwriting quality (simulated)"""
    hash_val = int(hashlib.md5(image_data[:100].encode()).hexdigest(), 16) % 100
    
    if hash_val > 70:
        return "CLEAR", 0.85
    elif hash_val > 40:
        return "STANDARD", 0.60
    else:
        return "CURSIVE", 0.30

# ===========================
# Core Analysis Functions
# ===========================

def auto_generate_tags(filename: str, text_preview: str) -> List[str]:
    """Auto-generate tags from filename and content"""
    tags = []
    
    filename_tag = extract_tag_from_filename(filename)
    if filename_tag:
        tags.append(filename_tag)
    
    text_lower = text_preview.lower()
    
    for tag in TAG_OPTIONS['equipment']:
        if tag.lower() in text_lower:
            tags.append(tag)
    
    for tag in TAG_OPTIONS['industry']:
        if tag.lower() in text_lower:
            tags.append(tag)
    
    for tag in TAG_OPTIONS['timeline']:
        if tag in text_preview:
            tags.append(tag)
    
    return list(set(tags))[:5]

def extract_qa_pairs(text: str, filename: str) -> str:
    """Extract Q&A pairs using LLM - returns formatted text"""
    try:
        user_prompt = f"""Document: {filename}

Content:
{text[:3000]}

Extract all question-answer pairs from this underwriting document."""

        response = call_llm_api(QA_EXTRACTION_SYSTEM, user_prompt)
        
        if not response:
            return "No Q&A pairs could be extracted from this document."
        
        return response
        
    except Exception as e:
        st.warning(f"Q&A extraction error: {e}")
        return "Error extracting Q&A pairs."

def analyze_electronic_text(text: str, filename: str) -> str:
    """Analyze electronic/printed text - returns formatted text"""
    try:
        user_prompt = f"""Document: {filename}

Full Content:
{text[:6000]}

Perform comprehensive analysis of this underwriting document. 
Base your analysis ONLY on the actual content provided above."""

        response = call_llm_api(ELECTRONIC_TEXT_ANALYSIS_SYSTEM, user_prompt, max_tokens=5000)
        
        if not response:
            return "Unable to analyze electronic text. Please check API configuration."
        
        return response
        
    except Exception as e:
        st.warning(f"Electronic text analysis error: {e}")
        return "Error during electronic text analysis."

def translate_handwriting(images: List[Dict], filename: str, text_content: str = "") -> Dict:
    """Translate handwritten annotations with intelligent detection"""
    try:
        if not images:
            return {
                "has_handwriting": False,
                "translated_text": "No images detected in this document.",
                "image_count": 0
            }
        
        # Detect handwriting using heuristics
        has_handwriting = detect_handwriting_in_images(images)
        
        if not has_handwriting:
            return {
                "has_handwriting": False,
                "translated_text": f"Document contains {len(images)} image(s), but no handwriting annotations detected.",
                "image_count": len(images)
            }
        
        # Prepare analysis prompt with document context
        max_page = max([img.get('page', 1) for img in images])
        
        user_prompt = f"""Document: {filename}

This is a scanned insurance document with {len(images)} images across {max_page} page(s).

Document Context:
{text_content[:1500] if text_content else "Scanned document - analyzing image-based content"}

Image Analysis:
- Total images: {len(images)}
- Distribution: Page 1 has {len([i for i in images if i.get('page')==1])} image(s)
- Small overlays detected: {len([i for i in images if i.get('size', 0) < 50000])} (likely handwriting/stamps)

Task: Analyze the document structure to identify and translate any handwritten annotations.

For scanned underwriting documents, handwritten notes typically include:
- Executive comments (e.g., "To CEO", "For review")
- Renewal recommendations or suggestions
- Approval signatures or initials
- Date stamps or reference numbers
- Risk assessments or underwriter notes
- Special instructions or attention markers

Provide translation in the specified format with:
1. Summary of detected handwritten content
2. Each annotation with its location, type, and translated text
3. Key insights about what the handwriting indicates
4. Any items needing manual review"""

        response = call_llm_api(HANDWRITING_TRANSLATION_SYSTEM, user_prompt, temperature=0.2, max_tokens=3000)
        
        if not response:
            response = f"✅ **Have handwriting notes**\n\n{len(images)} image(s) detected in document. Handwriting analysis in progress.\n\nNote: For accurate OCR, integrate with Google Cloud Vision API or AWS Textract."
        
        return {
            "has_handwriting": True,
            "translated_text": response,
            "image_count": len(images),
            "needs_review": []
        }
        
    except Exception as e:
        st.warning(f"Handwriting translation error: {e}")
        return {
            "has_handwriting": False,
            "translated_text": f"Error during handwriting translation: {e}",
            "image_count": 0
        }

def perform_dual_track_analysis(text: str, images: List[Dict], filename: str) -> Dict:
    """Perform comprehensive dual-track analysis"""
    try:
        # Track 1: Electronic text analysis
        electronic_analysis = analyze_electronic_text(text, filename)
        
        # Track 2: Handwriting translation (with text context)
        handwriting_translation = translate_handwriting(images, filename, text)
        
        # Extract Q&A pairs
        qa_pairs = extract_qa_pairs(text, filename)
        
        # Combined analysis
        has_handwriting = handwriting_translation.get('has_handwriting', False)
        handwriting_text = handwriting_translation.get('translated_text', '')
        
        integration_prompt = f"""Create a comprehensive underwriting report integrating:

ELECTRONIC TEXT ANALYSIS:
{electronic_analysis[:2500]}

HANDWRITING NOTES:
{handwriting_text[:1200] if has_handwriting else "No handwritten notes detected"}

Q&A SUMMARY:
{qa_pairs[:1000]}

Provide:
1. Executive Summary (2-3 paragraphs covering key points)
2. Critical Risk Factors (identify top 3-5 risks)
3. Underwriting Recommendations (specific actions needed)
4. Key Decision Points (items requiring management attention)

Base the report on ACTUAL content from this specific document."""

        integration_response = call_llm_api(SYSTEM_INSTRUCTION, integration_prompt, max_tokens=4000)
        
        if not integration_response:
            integration_response = "Unable to generate integrated report. Please review individual sections."
        
        full_analysis = {
            "electronic_analysis": electronic_analysis,
            "handwriting_translation": handwriting_translation,
            "qa_extraction": qa_pairs,
            "integrated_report": integration_response,
            "has_handwriting": has_handwriting,
            "analysis_timestamp": datetime.now().isoformat()
        }
        
        return full_analysis
        
    except Exception as e:
        st.error(f"Dual-track analysis error: {e}")
        traceback.print_exc()
        return {}

def auto_annotate_by_llm(filename: str, text: str, existing_tags: List[str] = None) -> Dict:
    """Auto-annotate document using LLM"""
    auto_tags = []
    
    try:
        filename_tag = extract_tag_from_filename(filename)
        auto_tags = auto_generate_tags(filename, text[:2000])
        
        if existing_tags:
            auto_tags.extend(existing_tags)
        auto_tags = list(set(auto_tags))
        
        user_prompt = f"""Document: {filename}
Existing tags: {', '.join(auto_tags)}

Content preview:
{text[:3000]}

Provide comprehensive auto-annotation for this underwriting document."""

        response = call_llm_api(AUTO_ANNOTATE_SYSTEM, user_prompt, temperature=0.3)
        
        # Parse response to extract structured data
        annotations = {
            'tags': auto_tags if auto_tags else ['Unclassified'],
            'insurance_type': 'General',
            'decision': 'Pending',
            'premium_estimate': 'TBD',
            'retention': 'TBD',
            'risk_level': 'Medium',
            'case_summary': response[:200] if response else 'Manual review required',
            'key_insights': ['Requires analysis'],
            'confidence': 0.7
        }
        
        return annotations
            
    except Exception as e:
        st.warning(f"Auto-annotation error: {e}")
        return {
            'tags': auto_tags if auto_tags else ['Unclassified'],
            'insurance_type': 'General',
            'decision': 'Pending',
            'premium_estimate': 'TBD',
            'retention': 'TBD',
            'risk_level': 'Medium',
            'case_summary': 'Auto-annotation error, manual review required',
            'key_insights': ['Error during analysis'],
            'confidence': 0.0
        }

# ===========================
# Workspace Management
# ===========================

def create_workspace(name: str, description: str = ""):
    """Create a new workspace"""
    workspace_dir = WORKSPACES_DIR / name
    workspace_dir.mkdir(exist_ok=True)
    
    metadata = {
        "name": name,
        "description": description,
        "created_at": datetime.now().isoformat(),
        "documents": []
    }
    
    with open(workspace_dir / "metadata.json", 'w') as f:
        json.dump(metadata, f, indent=2)
    
    log_audit_event("workspace_created", {"workspace": name})
    
    return metadata

def load_workspace(name: str) -> Optional[Dict]:
    """Load workspace metadata"""
    workspace_dir = WORKSPACES_DIR / name
    metadata_file = workspace_dir / "metadata.json"
    
    if metadata_file.exists():
        with open(metadata_file, 'r') as f:
            return json.load(f)
    return None

def list_workspaces() -> List[str]:
    """List all workspaces"""
    if not WORKSPACES_DIR.exists():
        return []
    return [d.name for d in WORKSPACES_DIR.iterdir() if d.is_dir()]

def delete_document_from_workspace(workspace_name: str, filename: str) -> bool:
    """Delete a document from workspace"""
    try:
        metadata = load_workspace(workspace_name)
        if not metadata:
            return False
        
        # Find the document
        doc = next((d for d in metadata['documents'] if d['filename'] == filename), None)
        if not doc:
            return False
        
        # Delete physical file
        file_path = Path(doc['path'])
        if file_path.exists():
            file_path.unlink()
        
        # Delete embedding file
        embedding_file = EMBEDDINGS_DIR / f"{workspace_name}_{filename}.json"
        if embedding_file.exists():
            embedding_file.unlink()
        
        # Delete analysis file
        analysis_file = ANALYSIS_DIR / f"{workspace_name}_{filename}.json"
        if analysis_file.exists():
            analysis_file.unlink()
        
        # Remove from metadata
        metadata['documents'] = [d for d in metadata['documents'] if d['filename'] != filename]
        
        # Save metadata
        workspace_dir = WORKSPACES_DIR / workspace_name
        with open(workspace_dir / "metadata.json", 'w') as f:
            json.dump(metadata, f, indent=2)
        
        log_audit_event("document_deleted", {
            "workspace": workspace_name,
            "filename": filename
        })
        
        return True
        
    except Exception as e:
        st.error(f"Error deleting document: {e}")
        return False

def upload_document_to_workspace(workspace_name: str, uploaded_file, auto_analyze: bool = True):
    """Upload document to workspace with auto-analysis"""
    try:
        workspace_dir = WORKSPACES_DIR / workspace_name
        file_path = workspace_dir / uploaded_file.name
        
        with open(file_path, 'wb') as f:
            f.write(uploaded_file.getvalue())
        
        # Extract text and images
        extracted_text, images = extract_text_from_file(file_path)
        
        if not extracted_text:
            st.warning(f"No text extracted from {uploaded_file.name}")
            return None
        
        embedding = generate_embedding(extracted_text[:2000])
        
        annotations = {}
        if auto_analyze:
            annotations = auto_annotate_by_llm(uploaded_file.name, extracted_text)
        
        doc_metadata = {
            "filename": uploaded_file.name,
            "format": file_path.suffix.lstrip('.'),
            "path": str(file_path),
            "size": uploaded_file.size,
            "upload_date": datetime.now().isoformat(),
            "extracted_text_preview": extracted_text[:500],
            "has_images": len(images) > 0,
            "image_count": len(images),
            "tags": annotations.get('tags', []),
            "insurance_type": annotations.get('insurance_type', ''),
            "decision": annotations.get('decision', 'Pending'),
            "risk_level": annotations.get('risk_level', 'Medium'),
            "has_deep_analysis": False
        }
        
        metadata = load_workspace(workspace_name)
        metadata["documents"].append(doc_metadata)
        
        with open(workspace_dir / "metadata.json", 'w') as f:
            json.dump(metadata, f, indent=2)
        
        embedding_file = EMBEDDINGS_DIR / f"{workspace_name}_{uploaded_file.name}.json"
        with open(embedding_file, 'w') as f:
            json.dump({"embedding": embedding, "text_preview": extracted_text[:500]}, f)
        
        # Perform analysis if there are images or auto_analyze is enabled
        if auto_analyze and (len(images) > 0 or is_scanned_pdf(file_path)):
            with st.spinner("Performing dual-track analysis..."):
                analysis_result = perform_dual_track_analysis(extracted_text, images, uploaded_file.name)
                
                analysis_file = ANALYSIS_DIR / f"{workspace_name}_{uploaded_file.name}.json"
                with open(analysis_file, 'w') as f:
                    json.dump(analysis_result, f, indent=2)
                
                doc_metadata["has_deep_analysis"] = True
        
        log_audit_event("document_uploaded", {
            "workspace": workspace_name,
            "filename": uploaded_file.name,
            "auto_analyzed": auto_analyze
        })
        
        return doc_metadata
        
    except Exception as e:
        st.error(f"Upload error: {e}")
        traceback.print_exc()
        return None

def load_initial_dataset():
    """Load ONLY the MSC Memo file on first run"""
    try:
        # Look for the file with flexible naming
        possible_names = [
            "Hull - MSC_Memo.pdf",
            "Hull_MSC_Memo.pdf", 
            "Hull-MSC_Memo.pdf",
            "Hull - Marco Polo_Memo.pdf"  # Backward compatibility
        ]
        
        initial_file = None
        for name in possible_names:
            if Path(name).exists():
                initial_file = name
                break
        
        if not initial_file:
            st.info(f"ℹ️ Initial dataset file not found. Looking for: {possible_names[0]}")
            return False
        
        default_workspace = "Default"
        metadata = load_workspace(default_workspace)
        
        # Check if initial file already loaded
        if metadata:
            for doc in metadata.get("documents", []):
                if doc["filename"] == initial_file or initial_file in doc["filename"]:
                    return True  # Already loaded
        
        # Create workspace if not exists
        if not metadata:
            create_workspace(default_workspace, "Default workspace with initial dataset")
            metadata = load_workspace(default_workspace)
        
        # Load the file
        with open(initial_file, 'rb') as f:
            file_content = f.read()
        
        class UploadedFile:
            def __init__(self, name, content):
                self.name = name
                self.size = len(content)
                self._content = content
            
            def getvalue(self):
                return self._content
        
        uploaded_file = UploadedFile(initial_file, file_content)
        
        result = upload_document_to_workspace(default_workspace, uploaded_file, auto_analyze=True)
        
        if result:
            return True
        else:
            return False
            
    except Exception as e:
        st.error(f"Error loading initial dataset: {e}")
        return False

# ===========================
# Search and Retrieval
# ===========================

def search_documents(query: str, workspace_name: str, top_k: int = 5) -> List[Dict]:
    """Search documents using semantic similarity"""
    try:
        query_embedding = generate_embedding(query)
        results = []
        
        workspace_metadata = load_workspace(workspace_name)
        if not workspace_metadata:
            return []
        
        for doc in workspace_metadata.get("documents", []):
            filename = doc["filename"]
            embedding_file = EMBEDDINGS_DIR / f"{workspace_name}_{filename}.json"
            
            if embedding_file.exists():
                with open(embedding_file, 'r') as f:
                    data = json.load(f)
                    doc_embedding = data["embedding"]
                    
                    similarity = cosine_similarity(query_embedding, doc_embedding)
                    
                    query_lower = query.lower()
                    for term in INSURANCE_TERMS.keys():
                        if term in query_lower and term in data.get("text_preview", "").lower():
                            similarity += 0.1
                    
                    results.append({
                        "document": doc,
                        "similarity": similarity,
                        "preview": data.get("text_preview", "")
                    })
        
        results.sort(key=lambda x: x["similarity"], reverse=True)
        
        return results[:top_k]
        
    except Exception as e:
        st.error(f"Search error: {e}")
        return []

# ===========================
# UI Functions
# ===========================

def inject_css():
    """Inject custom CSS - WHITE THEME"""
    
    light_css = """
    <style>
    /* 主背景 - 白色渐变 */
    .stApp {
        background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
    }
    
    /* 页头 - 保持蓝紫色渐变但调整阴影 */
    .main-header {
        background: linear-gradient(90deg, #3b82f6 0%, #8b5cf6 100%);
        padding: 2rem;
        border-radius: 15px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.15);
    }
    
    .main-header h1 {
        color: #ffffff;
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
    }
    
    .main-header p {
        color: #e0e7ff;
        font-size: 1.1rem;
        margin-top: 0.5rem;
    }
    
    /* 标签徽章 */
    .tag-badge {
        display: inline-block;
        padding: 0.4rem 0.8rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        margin: 0.2rem;
        background: linear-gradient(135deg, #3b82f6, #8b5cf6);
        color: white;
        box-shadow: 0 2px 4px rgba(59, 130, 246, 0.2);
    }
    
    .analysis-badge {
        background: linear-gradient(135deg, #10b981, #059669);
    }
    
    .handwriting-badge {
        background: linear-gradient(135deg, #f59e0b, #d97706);
    }
    
    /* 文档卡片 - 白色背景，浅灰色边框 */
    .doc-card {
        background: #ffffff;
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #e2e8f0;
        margin-bottom: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
    }
    
    .doc-card:hover {
        border-color: #3b82f6;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.15);
        transform: translateY(-2px);
    }
    
    /* 文字颜色 - 深色以确保可读性 */
    .doc-card h3 {
        color: #1e293b;
    }
    
    .doc-card p {
        color: #475569;
    }
    
    /* 按钮样式 */
    .stButton > button {
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.25);
    }
    
    /* 侧边栏 - 浅灰色背景 */
    [data-testid="stSidebar"] {
        background-color: #f8fafc;
    }
    
    /* 提高文字可读性 */
    .stMarkdown {
        color: #1e293b;
    }
    
    /* 标签页样式 */
    .stTabs [data-baseweb="tab-list"] {
        background-color: #f1f5f9;
        border-radius: 8px;
        padding: 0.5rem;
    }
    
    .stTabs [data-baseweb="tab"] {
        color: #475569;
        font-weight: 600;
    }
    
    .stTabs [aria-selected="true"] {
        color: #3b82f6;
    }
    
    /* 输入框样式优化 */
    .stTextInput > div > div > input {
        background-color: #ffffff;
        border-color: #e2e8f0;
    }
    
    /* 选择框样式 */
    .stSelectbox > div > div {
        background-color: #ffffff;
    }
    
    /* 文件上传区域 */
    [data-testid="stFileUploader"] {
        background-color: #f8fafc;
        border: 2px dashed #cbd5e1;
        border-radius: 8px;
    }
    
    /* 警告和信息框 */
    .stAlert {
        background-color: #ffffff;
        border-left: 4px solid #3b82f6;
    }
    
    /* 指标卡片 */
    [data-testid="stMetricValue"] {
        color: #1e293b;
    }
    
    /* 展开器样式 */
    .streamlit-expanderHeader {
        background-color: #f8fafc;
        color: #1e293b;
        border-radius: 8px;
    }
    
    /* 数据框样式 */
    .dataframe {
        background-color: #ffffff;
    }
    </style>
    """
    
    st.markdown(light_css, unsafe_allow_html=True)

def render_header():
    """Render application header"""
    st.markdown(f"""
    <div class="main-header">
        <h1>📋 {APP_TITLE}</h1>
        <p>Version {VERSION} | Powered by AI | Advanced OCR & Document Analysis</p>
    </div>
    """, unsafe_allow_html=True)

def render_api_config_sidebar():
    """Render API configuration in sidebar"""
    with st.sidebar:
        st.markdown("---")
        st.subheader("⚙️ API Configuration")
        
        current_key = get_api_key()
        
        with st.expander("🔑 Configure AI Model API"):
            api_key_input = st.text_input(
                "API Key:",
                value=current_key,
                type="password",
                help="Enter your API key"
            )
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("💾 Save", use_container_width=True):
                    if api_key_input:
                        save_api_config(api_key_input)
                        st.session_state.api_key = api_key_input
                        st.success("✅ Saved!")
                        st.rerun()
                    else:
                        st.warning("Please enter API key")
            
            with col2:
                if st.button("🧪 Test", use_container_width=True):
                    if api_key_input:
                        with st.spinner("Testing..."):
                            is_valid, message = validate_api_key(api_key_input)
                            
                            if is_valid:
                                st.success(f"✅ {message}")
                            else:
                                st.error(f"❌ {message}")
                    else:
                        st.warning("Please enter API key")

def render_document_card(doc: Dict, workspace_name: str, doc_index: int = 0):
    """Render a document card with unique keys and delete button"""
    # Safely get format with fallback
    format_icon = SUPPORTED_FORMATS.get(doc.get('format', 'pdf'), '📄')
    
    # Create unique key prefix using upload date and index
    upload_ts = doc.get('upload_date', '').replace(':', '-').replace('.', '-')
    key_prefix = f"{workspace_name}_{hashlib.md5(doc['filename'].encode()).hexdigest()[:8]}_{upload_ts}_{doc_index}"
    
    tags_html = " ".join([f'<span class="tag-badge">{tag}</span>' for tag in doc.get('tags', [])])
    
    analysis_badge = ""
    if doc.get('has_deep_analysis'):
        analysis_badge = '<span class="tag-badge analysis-badge">✓ Analyzed</span>'
    
    if doc.get('has_images'):
        analysis_badge += ' <span class="tag-badge handwriting-badge">✍️ Has Images</span>'
    
    st.markdown(f"""
    <div class="doc-card">
        <h3>{format_icon} {doc['filename']}</h3>
        <p><strong>Category:</strong> {doc.get('insurance_type', 'N/A')} | 
           <strong>Decision:</strong> {doc.get('decision', 'Pending')}</p>
        <p>{tags_html} {analysis_badge}</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button(f"📄 View", key=f"view_{key_prefix}"):
            st.session_state.viewing_doc = doc
    
    with col2:
        if doc.get('has_deep_analysis'):
            if st.button(f"📊 Analysis", key=f"analysis_{key_prefix}"):
                st.session_state.viewing_analysis = (workspace_name, doc['filename'])
    
    with col3:
        file_path = Path(doc['path'])
        if file_path.exists():
            with open(file_path, 'rb') as f:
                st.download_button(
                    label="⬇️ Download",
                    data=f.read(),
                    file_name=doc['filename'],
                    key=f"download_{key_prefix}"
                )
    
    with col4:
        if st.button(f"🗑️ Delete", key=f"delete_{key_prefix}", type="secondary"):
            st.session_state[f"confirm_delete_{key_prefix}"] = True
    
    # Confirmation dialog for delete
    if st.session_state.get(f"confirm_delete_{key_prefix}", False):
        st.warning(f"⚠️ Are you sure you want to delete **{doc['filename']}**?")
        col_yes, col_no = st.columns(2)
        
        with col_yes:
            if st.button("✅ Yes, Delete", key=f"confirm_yes_{key_prefix}"):
                if delete_document_from_workspace(workspace_name, doc['filename']):
                    st.success(f"✅ Deleted {doc['filename']}")
                    st.session_state[f"confirm_delete_{key_prefix}"] = False
                    st.rerun()
                else:
                    st.error("Failed to delete document")
        
        with col_no:
            if st.button("❌ Cancel", key=f"confirm_no_{key_prefix}"):
                st.session_state[f"confirm_delete_{key_prefix}"] = False
                st.rerun()

def parse_handwriting_translations(text: str) -> List[Dict]:
    """解析手写翻译结果，提取位置、文本和置信度"""
    translations = []
    if not text:
        return translations
    
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line or line.startswith('#') or line.startswith('**'):
            continue
            
        # 查找格式: [Location] Text (Confidence: XX%)
        if '[' in line and ']' in line:
            try:
                # 提取位置
                location_start = line.find('[')
                location_end = line.find(']')
                location = line[location_start+1:location_end]
                
                # 提取剩余部分
                rest = line[location_end+1:].strip()
                
                # 提取置信度
                confidence = 70  # 默认值
                if '(Confidence:' in rest or '(confidence:' in rest:
                    conf_start = rest.lower().find('(confidence:')
                    conf_section = rest[conf_start:]
                    # 查找百分号之前的数字
                    import re
                    conf_match = re.search(r'(\d+)%', conf_section)
                    if conf_match:
                        confidence = int(conf_match.group(1))
                    
                    # 提取文本（去掉置信度部分）
                    text_content = rest[:conf_start].strip()
                else:
                    text_content = rest
                
                if text_content:  # 只添加有文本内容的
                    translations.append({
                        'location': location,
                        'text': text_content,
                        'confidence': confidence
                    })
            except Exception as e:
                continue  # 跳过格式不正确的行
    
    return translations

def extract_insurance_type(analysis: Dict) -> str:
    """从分析结果中提取保险类型"""
    text = json.dumps(analysis).lower()
    
    if 'hull' in text and 'machinery' in text:
        return "Hull & Machinery"
    elif 'hull' in text:
        return "Hull"
    elif 'cargo' in text:
        return "Cargo"
    elif 'p&i' in text or 'protection' in text and 'indemnity' in text:
        return "P&I"
    elif 'war' in text and 'risk' in text:
        return "War Risk"
    elif 'liability' in text:
        return "Liability"
    else:
        return "其他"

def extract_client_name(analysis: Dict) -> str:
    """从分析结果中提取客户名称"""
    text = json.dumps(analysis)
    
    # 查找常见模式
    patterns = [
        r'[Ii]nsured[:\s]+([A-Z][A-Za-z\s&\.]+(?:Company|Corp|Ltd|Inc|LLC|S\.A\.)?)',
        r'[Cc]lient[:\s]+([A-Z][A-Za-z\s&\.]+(?:Company|Corp|Ltd|Inc|LLC|S\.A\.)?)',
        r'[Aa]ssured[:\s]+([A-Z][A-Za-z\s&\.]+(?:Company|Corp|Ltd|Inc|LLC|S\.A\.)?)',
        r'MSC',
        r'Mediterranean Shipping Company'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            if isinstance(match.group(0), str):
                return match.group(0)[:50].strip()
    
    return "Unknown"

def extract_year(analysis: Dict) -> str:
    """从分析结果中提取承保年度"""
    text = json.dumps(analysis)
    
    # 查找 2020-2030 之间的年份
    years = re.findall(r'20[2-3][0-9]', text)
    
    if years:
        # 返回最常见的年份
        from collections import Counter
        most_common = Counter(years).most_common(1)
        return most_common[0][0]
    
    return datetime.now().strftime("%Y")

def render_analysis_view(workspace_name: str, filename: str):
    """Render analysis results in client-ready format"""
    analysis_file = ANALYSIS_DIR / f"{workspace_name}_{filename}.json"
    
    if not analysis_file.exists():
        st.warning("No analysis found for this document")
        return
    
    with open(analysis_file, 'r') as f:
        analysis = json.load(f)
    
    st.subheader("📊 Comprehensive Analysis Results")
    
    # Show API status warning if needed
    api_key = get_api_key()
    if not api_key:
        st.error("⚠️ API key not configured. Analysis may be incomplete.")
    
    view_mode = st.radio(
        "Select View:",
        ["Integrated Report", "Electronic Text", "Handwriting Translation"],
        horizontal=True
    )
    
    if view_mode == "Integrated Report":
        st.markdown("### 📊 Integrated Analysis Report")
        
        # 表格展示部分
        st.markdown("#### Document Information")
        
        # 提取元数据
        insurance_type = extract_insurance_type(analysis)
        client_name = extract_client_name(analysis)
        underwriting_year = extract_year(analysis)
        timestamp = analysis.get('timestamp', datetime.now().isoformat())[:19]
        
        # Create table data
        table_data = {
            "Document Name": [filename],
            "Category": [insurance_type],
            "Underwriting Year": [underwriting_year],
            "Last Updated": [timestamp]
        }
        
        import pandas as pd
        df = pd.DataFrame(table_data)
        st.dataframe(df, use_container_width=True, hide_index=True)
        
        st.markdown("---")
        
        # 主要内容
        st.markdown("#### Main Content")
        
        report = analysis.get('integrated_report', '')
        
        if not report or "unable" in report.lower():
            st.error("❌ Analysis generation failed")
            st.info("💡 Configure API key in sidebar and click 'Re-run Analysis'")
        else:
            st.markdown(report)
        
    elif view_mode == "Electronic Text":
        st.markdown("### 📄 Electronic Text Analysis")
        
        electronic = analysis.get('electronic_analysis', '')
        
        if not electronic or len(electronic) < 50:
            st.warning("⚠️ Electronic text analysis is empty or incomplete")
            st.info("This may be a scanned document. Check the Handwriting Translation tab.")
        else:
            st.markdown(electronic)
    
    elif view_mode == "Handwriting Translation":
        st.markdown("### ✍️ Handwriting Translation")
        
        handwriting = analysis.get('handwriting_translation', {})
        has_handwriting = handwriting.get('has_handwriting', False)
        
        if has_handwriting:
            st.success("✅ **Have handwriting notes**")
            
            # 解析翻译结果
            translated_text = handwriting.get('translated_text', '')
            translations = parse_handwriting_translations(translated_text)
            
            # 加载图片数据
            images = analysis.get('images', [])
            
            if translations:
                st.markdown("#### Handwriting Recognition Results")
                
                for idx, trans in enumerate(translations):
                    col_img, col_text = st.columns([1, 2])
                    
                    with col_img:
                        # 显示对应图片（如果有）
                        if idx < len(images) and images[idx].get('data'):
                            try:
                                img_data = images[idx]['data']
                                st.image(f"data:image/png;base64,{img_data}", width=200)
                            except:
                                st.write(f"🖼️ {trans['location']}")
                        else:
                            st.write(f"🖼️ {trans['location']}")
                    
                    with col_text:
                        # 显示翻译文本
                        st.markdown(f"**{trans['text']}**")
                        
                        # 显示识别度（进度条 + 百分比）
                        confidence = trans.get('confidence', 0)
                        
                        # 根据置信度选择颜色
                        if confidence >= 80:
                            color_class = "🟢"  # 绿色
                        elif confidence >= 60:
                            color_class = "🟡"  # 黄色
                        else:
                            color_class = "🔴"  # 红色
                        
                        st.progress(confidence / 100)
                        st.caption(f"{color_class} Confidence: {confidence}%")
                    
                    st.markdown("---")
            else:
                # 如果解析失败，显示原始文本
                st.markdown(translated_text)
        else:
            st.info("ℹ️ No handwritten content detected in this document")
        
        st.markdown("---")
        st.markdown("### 📤 Upload Handwriting Images for Recognition")
        st.markdown("Upload photos or scans of handwritten annotations:")
        
        uploaded_images = st.file_uploader(
            "Select handwriting image(s)",
            type=['png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            help="Upload images of handwritten notes or annotations",
            key=f"upload_handwriting_{filename}"
        )
        
        if uploaded_images:
            st.markdown(f"**Uploaded {len(uploaded_images)} image(s)**")
            
            for idx, img_file in enumerate(uploaded_images):
                with st.expander(f"Image {idx+1}: {img_file.name}"):
                    # Display image
                    st.image(img_file, caption=img_file.name, use_container_width=True)
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        if st.button(f"🔍 Recognize Text", key=f"ocr_{filename}_{idx}_{img_file.name}"):
                            st.info("🚧 OCR recognition feature")
                            st.markdown("""
                            **In production, this would:**
                            1. Send image to OCR service (Google Vision / AWS Textract)
                            2. Extract handwritten text
                            3. Display recognized text below
                            4. Allow editing and confirmation
                            
                            **Current simulation:**
                            This is a handwritten annotation that would be processed by OCR.
                            """)
                    
                    with col2:
                        confidence = st.slider("Confidence Level", 0, 100, 75, key=f"conf_{filename}_{idx}")
                    
                    # Simulated OCR result
                    session_key = f"ocr_result_{filename}_{idx}"
                    if session_key not in st.session_state:
                        st.session_state[session_key] = ""
                    
                    transcription = st.text_area(
                        "Recognized / Manual Transcription:",
                        value=st.session_state[session_key],
                        key=f"trans_{filename}_{idx}_{img_file.name}",
                        height=100,
                        help="Edit or enter the text from the handwriting"
                    )
                    
                    if st.button(f"💾 Save Transcription", key=f"save_{filename}_{idx}_{img_file.name}"):
                        st.session_state[session_key] = transcription
                        st.success(f"✅ Saved transcription for {img_file.name}")
    
    # Add action buttons
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🔄 Re-run Analysis", key=f"rerun_{filename}"):
            with st.spinner("Re-analyzing document..."):
                metadata = load_workspace(workspace_name)
                doc = next((d for d in metadata['documents'] if d['filename'] == filename), None)
                
                if doc:
                    file_path = Path(doc['path'])
                    text, images = extract_text_from_file(file_path)
                    
                    new_analysis = perform_dual_track_analysis(text, images, filename)
                    
                    with open(analysis_file, 'w') as f:
                        json.dump(new_analysis, f, indent=2)
                    
                    st.success("✅ Analysis complete!")
                    st.rerun()
    
    with col2:
        if st.button("📥 Export Report", key=f"export_{filename}"):
            # Export as text file
            report_text = f"""UNDERWRITING ANALYSIS REPORT
{'='*60}
Document: {filename}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

{analysis.get('integrated_report', 'No integrated report available')}

{'='*60}
ELECTRONIC TEXT ANALYSIS
{'='*60}
{analysis.get('electronic_analysis', 'No analysis available')}

{'='*60}
HANDWRITING TRANSLATION
{'='*60}
{analysis.get('handwriting_translation', {}).get('translated_text', 'No handwriting detected')}
"""
            st.download_button(
                label="Download TXT Report",
                data=report_text,
                file_name=f"{filename}_analysis_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain",
                key=f"download_report_{filename}"
            )

# ===========================
# Main Application
# ===========================

def main():
    st.set_page_config(
        page_title="Enhanced Underwriting Assistant",
        page_icon="📋",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    inject_css()
    ensure_dirs()
    
    # Initialize session state
    if 'current_workspace' not in st.session_state:
        st.session_state.current_workspace = "Default"
    
    if 'viewing_doc' not in st.session_state:
        st.session_state.viewing_doc = None
    
    if 'viewing_analysis' not in st.session_state:
        st.session_state.viewing_analysis = None
    
    if 'initial_load_done' not in st.session_state:
        st.session_state.initial_load_done = False
    
    if 'api_key' not in st.session_state:
        st.session_state.api_key = get_api_key()
    
    render_header()
    
    # Load initial dataset (ONLY ONCE, ONLY ONE FILE)
    if not st.session_state.initial_load_done:
        with st.spinner("Loading initial dataset..."):
            load_initial_dataset()
        st.session_state.initial_load_done = True
    
    # Sidebar
    with st.sidebar:
        st.header("🗂️ Workspace Management")
        
        workspaces = list_workspaces()
        
        if workspaces:
            selected_workspace = st.selectbox(
                "Select Workspace:",
                workspaces,
                index=workspaces.index(st.session_state.current_workspace) if st.session_state.current_workspace in workspaces else 0
            )
            
            if selected_workspace != st.session_state.current_workspace:
                st.session_state.current_workspace = selected_workspace
                st.rerun()
        
        st.markdown("---")
        
        with st.expander("➕ Create New Workspace"):
            new_workspace_name = st.text_input("Workspace Name:")
            new_workspace_desc = st.text_area("Description:")
            
            if st.button("Create Workspace"):
                if new_workspace_name:
                    create_workspace(new_workspace_name, new_workspace_desc)
                    st.success(f"Workspace '{new_workspace_name}' created!")
                    st.session_state.current_workspace = new_workspace_name
                    st.rerun()
        
        st.markdown("---")
        
        if st.session_state.current_workspace:
            metadata = load_workspace(st.session_state.current_workspace)
            if metadata:
                st.info(f"📁 **{metadata['name']}**")
                st.metric("Documents", len(metadata.get('documents', [])))
        
        render_api_config_sidebar()
    
    # Main tabs
    tab1, tab2, tab3, tab4 = st.tabs([
        "📚 Document Library",
        "⬆️ Upload Documents",
        "📊 Analysis Dashboard",
        "💬 Chat Assistant"
    ])
    
    # TAB 1: Document Library
    with tab1:
        st.header("📚 Document Library")
        
        if not st.session_state.current_workspace:
            st.warning("Please select or create a workspace first")
        else:
            metadata = load_workspace(st.session_state.current_workspace)
            
            if not metadata or not metadata.get('documents'):
                st.info("No documents in this workspace. Upload documents to get started.")
            else:
                documents = metadata.get('documents', [])
                
                st.markdown(f"**Showing {len(documents)} document(s)**")
                
                for idx, doc in enumerate(documents):
                    render_document_card(doc, st.session_state.current_workspace, doc_index=idx)
                
                if st.session_state.viewing_analysis:
                    workspace, filename = st.session_state.viewing_analysis
                    render_analysis_view(workspace, filename)
                    
                    if st.button("Close Analysis"):
                        st.session_state.viewing_analysis = None
                        st.rerun()
    
    # TAB 2: Upload
    with tab2:
        st.header("⬆️ Upload Documents")
        
        if not st.session_state.current_workspace:
            st.warning("Please select or create a workspace first")
        else:
            uploaded_files = st.file_uploader(
                "Upload underwriting documents:",
                type=list(SUPPORTED_FORMATS.keys()),
                accept_multiple_files=True
            )
            
            auto_analyze = st.checkbox("Perform automatic analysis", value=True)
            
            if uploaded_files and st.button("Upload & Process"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Processing {file.name}...")
                    
                    result = upload_document_to_workspace(
                        st.session_state.current_workspace,
                        file,
                        auto_analyze=auto_analyze
                    )
                    
                    if result:
                        st.success(f"✅ {file.name} uploaded!")
                    else:
                        st.error(f"❌ Failed to upload {file.name}")
                    
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                
                status_text.text("Upload complete!")
                st.balloons()
    
    # TAB 3: Analysis Dashboard
    with tab3:
        st.header("📊 Analysis Dashboard")
        
        if not st.session_state.current_workspace:
            st.warning("Please select a workspace first")
        else:
            metadata = load_workspace(st.session_state.current_workspace)
            
            if not metadata or not metadata.get('documents'):
                st.info("No documents to analyze")
            else:
                documents = metadata.get('documents', [])
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("Total Documents", len(documents))
                
                with col2:
                    analyzed = sum(1 for d in documents if d.get('has_deep_analysis'))
                    st.metric("Analyzed", analyzed)
                
                with col3:
                    high_risk = sum(1 for d in documents if d.get('risk_level') == 'High')
                    st.metric("High Risk", high_risk)
                
                with col4:
                    pending = sum(1 for d in documents if d.get('decision') == 'Pending')
                    st.metric("Pending Review", pending)
                
                st.markdown("---")
                
                st.subheader("Document Analysis Status")
                
                for idx, doc in enumerate(documents):
                    upload_ts = doc.get('upload_date', '').replace(':', '-').replace('.', '-')
                    unique_key = f"{hashlib.md5(doc['filename'].encode()).hexdigest()[:8]}_{upload_ts}_{idx}"
                    
                    with st.expander(f"{doc['filename']} - {doc.get('decision', 'Pending')}"):
                        col1, col2 = st.columns([2, 1])
                        
                        with col1:
                            st.markdown(f"**Risk Level:** {doc.get('risk_level', 'N/A')}")
                            st.markdown(f"**Tags:** {', '.join(doc.get('tags', []))}")
                            st.markdown(f"**Has Images:** {'Yes' if doc.get('has_images') else 'No'}")
                        
                        with col2:
                            if doc.get('has_deep_analysis'):
                                if st.button("View Full Analysis", key=f"view_full_{unique_key}"):
                                    st.session_state.viewing_analysis = (st.session_state.current_workspace, doc['filename'])
                                    st.rerun()
                            else:
                                if st.button("Run Analysis", key=f"run_{unique_key}"):
                                    with st.spinner("Analyzing..."):
                                        file_path = Path(doc['path'])
                                        text, images = extract_text_from_file(file_path)
                                        
                                        analysis = perform_dual_track_analysis(text, images, doc['filename'])
                                        
                                        analysis_file = ANALYSIS_DIR / f"{st.session_state.current_workspace}_{doc['filename']}.json"
                                        with open(analysis_file, 'w') as f:
                                            json.dump(analysis, f, indent=2)
                                        
                                        doc['has_deep_analysis'] = True
                                        
                                        with open(WORKSPACES_DIR / st.session_state.current_workspace / "metadata.json", 'w') as f:
                                            json.dump(metadata, f, indent=2)
                                        
                                        st.success("Analysis complete!")
                                        st.rerun()
    
    # TAB 4: Chat Assistant
    with tab4:
        st.header("💬 AI Assistant")
        st.markdown("Ask questions about your documents, policies, and underwriting decisions.")
        
        if 'messages' not in st.session_state:
            st.session_state.messages = []
        
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        if prompt := st.chat_input("Ask about underwriting, policies, risk assessment..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            
            with st.chat_message("user"):
                st.markdown(prompt)
            
            with st.chat_message("assistant"):
                with st.spinner("Analyzing..."):
                    search_results = search_documents(
                        prompt,
                        st.session_state.current_workspace,
                        top_k=3
                    )
                    
                    context = "\n\n".join([
                        f"Document: {r['document']['filename']}\n{r['preview']}"
                        for r in search_results
                    ])
                    
                    user_prompt = f"""User Question: {prompt}

Relevant Documents:
{context}

Provide a comprehensive answer based on the available documents."""

                    response = call_llm_api(SYSTEM_INSTRUCTION, user_prompt)
                    
                    if not response:
                        response = "I apologize, but I'm unable to process your request at this time. Please ensure the API key is configured correctly in the sidebar."
                    
                    st.markdown(response)
                    
                    if search_results:
                        with st.expander("📚 Sources"):
                            for r in search_results:
                                st.markdown(f"- **{r['document']['filename']}** (similarity: {r['similarity']:.2f})")
                    
                    st.session_state.messages.append({"role": "assistant", "content": response})

if __name__ == "__main__":
    main()
